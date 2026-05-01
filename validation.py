import os
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document

# ======================================
# CONFIG
# ======================================


 

llm = ChatOpenAI(
    model="gpt-3.5-turbo",
    temperature=0
)

# ======================================
# SAVE FUNCTION
# ======================================
def save_to_docx(result, filename="llm_validation_result.docx"):
    # Note: This relies on the 'python-docx' package, NOT 'docx'.
    doc = Document()
    doc.add_heading("LLM Validation Result", 0)
    doc.add_paragraph(result)
    doc.save(filename)
    print(f"\n[Success] Result saved to {filename}")

# ======================================
# LLM JUDGE FUNCTION
# ======================================
def llm_judge(resume_text, predicted_role):

    template = """
    You are an expert career advisor.

    Resume:
    {resume_text}

    Predicted Role: {predicted_role}

    Is this role appropriate for the candidate based on their resume?

    Answer strictly in this format:
    Verdict: YES or NO
    Reason: <short reason>
    """

    prompt = PromptTemplate.from_template(template)
    chain = prompt | llm | StrOutputParser()

    try:
        response_text = chain.invoke({
            "resume_text": resume_text,
            "predicted_role": predicted_role
        })

        # FIX 2: Safer verdict extraction using the "in" operator
        # This handles cases where the LLM responds with "**Verdict:** YES"
        lines = response_text.upper().split("\n")
        verdict_line = next((l for l in lines if "VERDICT" in l), "")
        
        if "YES" in verdict_line:
            verdict = 1
        elif "NO" in verdict_line:
            verdict = 0
        else:
            verdict = -1 # Fallback if the LLM ignores formatting instructions

        # Display result
        print("\n===== LLM VALIDATION RESULT =====")
        print(f"Verdict (binary): {verdict}")
        print(response_text)

        # Save to docx
        save_to_docx(response_text)

        return verdict, response_text

    except Exception as e:
        print(f"\n[Error] Error during LLM validation: {str(e)}")
        return 0, str(e)

# ======================================
# TEST RUN
# ======================================
if __name__ == "__main__":
    sample_resume = """
    John Doe
    Skills: Python, Machine Learning, Data Analysis, TensorFlow, SQL
    Experience: 2 years as a Data Analyst at ABC Corp
    Education: B.Tech Computer Science
    """

    sample_role = "Data Scientist"

    verdict, explanation = llm_judge(sample_resume, sample_role)
    
    # Translate binary verdict back to readable text
    if verdict == 1:
        verdict_str = "Appropriate"
    elif verdict == 0:
        verdict_str = "Not Appropriate"
    else:
        verdict_str = "Parsing Failed (Check LLM Output)"
        
    print(f"\nFinal Verdict: {verdict_str}")
    print(f"Explanation: {explanation}")