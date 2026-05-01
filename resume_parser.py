# resume_parser.py
# Extracts text from PDF and DOCX resumes

import fitz         # pymupdf for PDF
from docx import Document
import io

def extract_text_from_pdf(file_bytes):
    """
    Extract all text from PDF file
    using PyMuPDF
    """
    try:
        pdf_document = fitz.open(
            stream = file_bytes,
            filetype = "pdf"
        )

        full_text = ""

        for page_num in range(
                len(pdf_document)):
            page = pdf_document.load_page(page_num)
            full_text += page.get_text()

        pdf_document.close()

        if not full_text.strip():
            return None, "No text found in PDF. "\
                        "It may be a scanned image."

        return full_text, None

    except Exception as e:
        return None, f"PDF reading error: {str(e)}"


def extract_text_from_docx(file_bytes):
    """
    Extract all text from DOCX file
    using python-docx
    """
    try:
        doc = Document(io.BytesIO(file_bytes))

        full_text = ""

        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + "\n"

        if not full_text.strip():
            return None, "No text found in DOCX file."

        return full_text, None

    except Exception as e:
        return None, f"DOCX reading error: {str(e)}"


def extract_text_from_file(uploaded_file):
    """
    Main function — detects file type
    and extracts text accordingly
    """
    file_bytes = uploaded_file.read()
    file_name  = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        text, error = extract_text_from_pdf(
                        file_bytes)

    elif file_name.endswith(".docx"):
        text, error = extract_text_from_docx(
                        file_bytes)

    elif file_name.endswith(".txt"):
        try:
            text  = file_bytes.decode("utf-8")
            error = None
        except Exception as e:
            text  = None
            error = f"TXT reading error: {str(e)}"

    else:
        text  = None
        error = "Unsupported file type. "\
                "Please upload PDF, DOCX or TXT."

    return text, error